import os
import uuid
import zipfile
import shutil
from flask import Flask, request, jsonify, render_template, send_from_directory, send_file
from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

app = Flask(__name__)

# CONFIGURATION
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
REPORT_FOLDER = os.path.join(BASE_DIR, 'reports')

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['REPORT_FOLDER'] = REPORT_FOLDER

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(REPORT_FOLDER, exist_ok=True)


# ----------------------------
# WORD DOCUMENT STYLING HELPERS
# ----------------------------

def set_cell_background(cell, hex_color):
    """Sets the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets the padding (margins) of a table cell in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{edge}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, **kwargs):
    """
    Sets borders for individual cell sides.
    Options: top, bottom, left, right, insideH, insideV
    Format: border_name={'val': 'single', 'sz': 4, 'color': 'HEXCOLOR'}
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), edge_data.get('val', 'single'))
            border.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), edge_data.get('color', 'auto'))
            tcBorders.append(border)
        else:
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
            
    tcPr.append(tcBorders)

def add_page_number(run):
    """Adds a dynamic page number field to a footer run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def style_paragraph(p, text="", font_name="Arial", font_size=11, bold=False, italic=False, color_rgb=None, align=0):
    """Applies styles to a paragraph and its run."""
    p.alignment = align
    if text:
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic
        if color_rgb:
            run.font.color.rgb = color_rgb
        return run
    return None

def find_file_case_insensitive(folder, base_name):
    """Finds a file in the folder matching the base_name, case-insensitively, with common extensions."""
    if not base_name:
        return None
    
    base_name = str(base_name).strip()
    # Direct check
    for ext in ['.jpg', '.jpeg', '.png', '.JPG', '.PNG', '.JPEG']:
        full_name = f"{base_name}{ext}"
        if os.path.exists(os.path.join(folder, full_name)):
            return full_name
            
    # List directory fallback
    try:
        files = os.listdir(folder)
    except FileNotFoundError:
        return None
        
    for f in files:
        name, ext = os.path.splitext(f)
        if name.lower().strip() == base_name.lower():
            return f
            
    return None


# ----------------------------
# FLASK ROUTING
# ----------------------------

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/uploads/<session_id>/<filename>')
def serve_upload(session_id, filename):
    session_images_dir = os.path.join(app.config['UPLOAD_FOLDER'], session_id, 'images')
    return send_from_directory(session_images_dir, filename)

@app.route('/api/upload', methods=['POST'])
def upload_files():
    session_id = request.form.get('session_id')
    if not session_id:
        session_id = str(uuid.uuid4())
        
    session_dir = os.path.join(app.config['UPLOAD_FOLDER'], session_id)
    excel_dir = os.path.join(session_dir, 'excel')
    images_dir = os.path.join(session_dir, 'images')
    
    os.makedirs(excel_dir, exist_ok=True)
    os.makedirs(images_dir, exist_ok=True)
    
    excel_file = request.files.get('excel_file')
    if excel_file:
        excel_path = os.path.join(excel_dir, 'Thermography_Master_Mapping.xlsx')
        excel_file.save(excel_path)
        
    # Process images (both files and zip archives)
    uploaded_files = request.files.getlist('images')
    for f in uploaded_files:
        if not f.filename:
            continue
            
        filename = f.filename
        filepath = os.path.join(images_dir, filename)
        f.save(filepath)
        
        # If it's a zip file, extract it
        if filename.lower().endswith('.zip'):
            try:
                with zipfile.ZipFile(filepath, 'r') as zip_ref:
                    # Extract directly to the images_dir
                    # We flatten any folder structures to make matching simple
                    for member in zip_ref.infolist():
                        if member.is_dir():
                            continue
                        # Get filename without path
                        filename_only = os.path.basename(member.filename)
                        if not filename_only:
                            continue
                        source = zip_ref.open(member)
                        target = open(os.path.join(images_dir, filename_only), "wb")
                        with source, target:
                            shutil.copyfileobj(source, target)
                # Remove the zip file after extraction
                os.remove(filepath)
            except Exception as e:
                return jsonify({'error': f'Failed to extract zip file: {str(e)}'}), 400

    # Parse Excel Sheet if it exists
    excel_path = os.path.join(excel_dir, 'Thermography_Master_Mapping.xlsx')
    if not os.path.exists(excel_path):
        return jsonify({
            'session_id': session_id,
            'locations': [],
            'client_name': 'D-MART',
            'site_name': 'KANCHARLAPALEM, VISAKHAPATNAM',
            'date': '11-05-2026',
            'message': 'Excel file missing. Please upload the mapping Excel file.'
        })

    try:
        wb = load_workbook(excel_path)
        ws = wb.active
        
        locations = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not row or row[0] is None:
                continue
                
            sno = row[0]
            location_name = row[1]
            r_img = row[2]
            y_img = row[3]
            b_img = row[4]
            n_img = row[5]
            visual_img = row[6]
            
            # Find corresponding file names on disk
            r_file = find_file_case_insensitive(images_dir, r_img)
            y_file = find_file_case_insensitive(images_dir, y_img)
            b_file = find_file_case_insensitive(images_dir, b_img)
            n_file = find_file_case_insensitive(images_dir, n_img)
            visual_file = find_file_case_insensitive(images_dir, visual_img)
            
            locations.append({
                'sno': sno,
                'location': location_name,
                'r_img': r_img,
                'y_img': y_img,
                'b_img': b_img,
                'n_img': n_img,
                'visual_img': visual_img,
                'r_file': r_file,
                'y_file': y_file,
                'b_file': b_file,
                'n_file': n_file,
                'visual_file': visual_file,
                # Default empty temperature inputs
                'r_temp': '',
                'y_temp': '',
                'b_temp': '',
                'n_temp': '',
                'observation': 'Temperature found to be Normal.',
                'remarks': 'Working in Normal Condition.'
            })
            
        return jsonify({
            'session_id': session_id,
            'locations': locations,
            'client_name': 'D-MART',
            'site_name': 'KANCHARLAPALEM, VISAKHAPATNAM',
            'date': '11-05-2026'
        })
    except Exception as e:
        return jsonify({'error': f'Error parsing Excel sheet: {str(e)}'}), 500


@app.route('/api/generate', methods=['POST'])
def generate_report():
    data = request.json
    if not data:
        return jsonify({'error': 'No data provided'}), 400
        
    session_id = data.get('session_id')
    client_name = data.get('client_name', 'D-MART').upper()
    site_name = data.get('site_name', 'KANCHARLAPALEM, VISAKHAPATNAM').upper()
    date_str = data.get('date', '11-05-2026')
    locations = data.get('locations', [])
    
    if not session_id:
        return jsonify({'error': 'Missing session_id'}), 400
        
    images_dir = os.path.join(app.config['UPLOAD_FOLDER'], session_id, 'images')
    session_report_dir = os.path.join(app.config['REPORT_FOLDER'], session_id)
    os.makedirs(session_report_dir, exist_ok=True)
    
    doc = Document()
    
    # ----------------------------
    # PAGE LAYOUT & MARGINS
    # ----------------------------
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
        # Configure Header & Footer
        section.different_first_page_header_footer = True
        
        # Secondary Header
        header = section.header
        header_p = header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header_p.add_run(f"ELECTRICAL THERMOGRAPHY SURVEY REPORT  |  {client_name}")
        header_run.font.name = 'Arial'
        header_run.font.size = Pt(8.5)
        header_run.font.color.rgb = RGBColor(100, 116, 139) # Slate color
        
        # Secondary Footer
        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.text = ""
        footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Left-aligned confidential text
        conf_run = footer_p.add_run("CONFIDENTIAL - ")
        conf_run.font.name = 'Arial'
        conf_run.font.size = Pt(8.5)
        conf_run.font.color.rgb = RGBColor(148, 163, 184)
        
        # Dynamic page number
        page_run = footer_p.add_run("Page ")
        page_run.font.name = 'Arial'
        page_run.font.size = Pt(8.5)
        page_run.font.color.rgb = RGBColor(100, 116, 139)
        add_page_number(page_run)

    # ----------------------------
    # COVER PAGE
    # ----------------------------
    # Create spacing from top
    for _ in range(3):
        doc.add_paragraph()
        
    # Title bar table
    title_table = doc.add_table(rows=1, cols=1)
    title_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_table.autofit = False
    title_table.columns[0].width = Inches(6.7)
    
    cell = title_table.cell(0, 0)
    set_cell_background(cell, "1E3A8A")  # Deep Navy Blue
    set_cell_margins(cell, top=500, bottom=500, left=400, right=400)
    
    # Title text
    p_title = cell.paragraphs[0]
    style_paragraph(p_title, "ELECTRICAL THERMOGRAPHY\nSURVEY REPORT", font_name="Arial", font_size=26, bold=True, color_rgb=RGBColor(255, 255, 255), align=1)
    
    # Subtitle
    doc.add_paragraph()
    p_sub = doc.add_paragraph()
    style_paragraph(p_sub, "Infrared Inspection & Thermal Analysis Report", font_name="Arial", font_size=14, italic=True, color_rgb=RGBColor(71, 85, 105), align=1)
    
    for _ in range(6):
        doc.add_paragraph()
        
    # Metadata Box (bordered panel)
    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_table.autofit = False
    meta_table.columns[0].width = Inches(2.2)
    meta_table.columns[1].width = Inches(4.5)
    
    labels = ["CLIENT NAME", "SITE LOCATION", "SURVEY DATE"]
    values = [client_name, site_name, date_str]
    
    border_style = {'val': 'single', 'sz': 4, 'color': 'E2E8F0'}
    
    for i in range(3):
        c0 = meta_table.cell(i, 0)
        c1 = meta_table.cell(i, 1)
        
        set_cell_background(c0, "F8FAFC")
        set_cell_margins(c0, top=120, bottom=120, left=150, right=150)
        set_cell_margins(c1, top=120, bottom=120, left=150, right=150)
        
        # Apply subtle borders
        set_cell_borders(c0, top=border_style, bottom=border_style, left=border_style, right=border_style)
        set_cell_borders(c1, top=border_style, bottom=border_style, left=border_style, right=border_style)
        
        p0 = c0.paragraphs[0]
        style_paragraph(p0, labels[i], font_name="Arial", font_size=9.5, bold=True, color_rgb=RGBColor(71, 85, 105))
        
        p1 = c1.paragraphs[0]
        style_paragraph(p1, values[i], font_name="Arial", font_size=10, bold=True, color_rgb=RGBColor(30, 41, 59))
        
    # Prepare for content
    doc.add_page_break()

    # ----------------------------
    # PROCESS LOCATIONS
    # ----------------------------
    for index, loc in enumerate(locations):
        sno = loc.get('sno', index + 1)
        loc_name = loc.get('location', 'Unknown Location').upper()
        
        # Title of Location
        p_loc = doc.add_paragraph()
        style_paragraph(p_loc, f"{sno}. {loc_name}", font_name="Arial", font_size=15, bold=True, color_rgb=RGBColor(30, 41, 59))
        
        # Add subtle horizontal line below header
        hr_table = doc.add_table(rows=1, cols=1)
        hr_table.autofit = False
        hr_table.columns[0].width = Inches(6.7)
        hr_cell = hr_table.cell(0, 0)
        set_cell_borders(hr_cell, bottom={'val': 'single', 'sz': 8, 'color': 'CBD5E1'})
        doc.add_paragraph()
        
        # ----------------------------
        # IMAGES LAYOUT (Side-by-side Table)
        # ----------------------------
        layout_table = doc.add_table(rows=1, cols=2)
        layout_table.autofit = False
        layout_table.columns[0].width = Inches(3.2)
        layout_table.columns[1].width = Inches(3.5)
        
        # Left Cell - Visual Image
        v_cell = layout_table.cell(0, 0)
        set_cell_margins(v_cell, top=0, bottom=0, left=0, right=100)
        set_cell_borders(v_cell) # Borderless
        
        p_v_label = v_cell.paragraphs[0]
        p_v_label.paragraph_format.space_after = Pt(4)
        style_paragraph(p_v_label, "Visual Reference:", font_name="Arial", font_size=9.5, bold=True, color_rgb=RGBColor(71, 85, 105))
        
        v_filename = loc.get('visual_file')
        v_img_path = os.path.join(images_dir, v_filename) if v_filename else ""
        
        if v_img_path and os.path.exists(v_img_path):
            p_v_img = v_cell.add_paragraph()
            p_v_img.add_run().add_picture(v_img_path, width=Inches(3.1))
        else:
            # Render visual placeholder
            placeholder_table = v_cell.add_table(rows=1, cols=1)
            placeholder_table.autofit = False
            placeholder_table.columns[0].width = Inches(3.1)
            ph_cell = placeholder_table.cell(0, 0)
            set_cell_background(ph_cell, "F1F5F9")
            set_cell_borders(ph_cell, top={'val': 'dashed', 'sz': 4, 'color': '94A3B8'},
                                      bottom={'val': 'dashed', 'sz': 4, 'color': '94A3B8'},
                                      left={'val': 'dashed', 'sz': 4, 'color': '94A3B8'},
                                      right={'val': 'dashed', 'sz': 4, 'color': '94A3B8'})
            set_cell_margins(ph_cell, top=1000, bottom=1000, left=100, right=100)
            p_ph = ph_cell.paragraphs[0]
            style_paragraph(p_ph, "[ Visual Reference Image ]\nNot Uploaded", font_name="Arial", font_size=9, color_rgb=RGBColor(148, 163, 184), align=1)
            v_cell.add_paragraph()
            
        # Right Cell - Thermal Images 2x2 Grid
        t_cell = layout_table.cell(0, 1)
        set_cell_margins(t_cell, top=0, bottom=0, left=100, right=0)
        set_cell_borders(t_cell) # Borderless
        
        p_t_label = t_cell.paragraphs[0]
        p_t_label.paragraph_format.space_after = Pt(4)
        style_paragraph(p_t_label, "Infrared Thermal Images:", font_name="Arial", font_size=9.5, bold=True, color_rgb=RGBColor(71, 85, 105))
        
        # Sub-table inside Right Cell for 2x2 grid
        grid_table = t_cell.add_table(rows=2, cols=2)
        grid_table.autofit = False
        grid_table.columns[0].width = Inches(1.68)
        grid_table.columns[1].width = Inches(1.68)
        
        phases = [
            {'name': 'R Phase', 'key': 'r_file'},
            {'name': 'Y Phase', 'key': 'y_file'},
            {'name': 'B Phase', 'key': 'b_file'},
            {'name': 'Neutral', 'key': 'n_file'}
        ]
        
        grid_positions = [(0, 0), (0, 1), (1, 0), (1, 1)]
        
        for idx, phase in enumerate(phases):
            r_idx, c_idx = grid_positions[idx]
            cell_ph = grid_table.cell(r_idx, c_idx)
            set_cell_margins(cell_ph, top=40, bottom=40, left=40, right=40)
            set_cell_borders(cell_ph) # Borderless
            
            p_name = cell_ph.paragraphs[0]
            p_name.paragraph_format.space_after = Pt(2)
            style_paragraph(p_name, phase['name'], font_name="Arial", font_size=8.5, bold=True, color_rgb=RGBColor(100, 116, 139))
            
            img_filename = loc.get(phase['key'])
            img_path = os.path.join(images_dir, img_filename) if img_filename else ""
            
            if img_path and os.path.exists(img_path):
                p_img = cell_ph.add_paragraph()
                p_img.add_run().add_picture(img_path, width=Inches(1.58))
            else:
                # Render thermal placeholder
                sub_ph_table = cell_ph.add_table(rows=1, cols=1)
                sub_ph_table.autofit = False
                sub_ph_table.columns[0].width = Inches(1.58)
                sub_ph_cell = sub_ph_table.cell(0, 0)
                set_cell_background(sub_ph_cell, "F8FAFC")
                set_cell_borders(sub_ph_cell, top={'val': 'dashed', 'sz': 4, 'color': 'CBD5E1'},
                                          bottom={'val': 'dashed', 'sz': 4, 'color': 'CBD5E1'},
                                          left={'val': 'dashed', 'sz': 4, 'color': 'CBD5E1'},
                                          right={'val': 'dashed', 'sz': 4, 'color': 'CBD5E1'})
                set_cell_margins(sub_ph_cell, top=450, bottom=450, left=50, right=50)
                p_sub_ph = sub_ph_cell.paragraphs[0]
                style_paragraph(p_sub_ph, "No Image", font_name="Arial", font_size=8, color_rgb=RGBColor(148, 163, 184), align=1)
                cell_ph.add_paragraph()
                
        doc.add_paragraph()
        
        # ----------------------------
        # TEMPERATURE MEASUREMENT TABLE
        # ----------------------------
        temp_table = doc.add_table(rows=5, cols=3)
        temp_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        temp_table.autofit = False
        temp_table.columns[0].width = Inches(2.2)
        temp_table.columns[1].width = Inches(2.2)
        temp_table.columns[2].width = Inches(2.3)
        
        # Header Row
        headers = ["PHASE", "MEASURED TEMPERATURE (°C)", "STATUS / RECOMMENDATION"]
        header_colors = ["1E3A8A", "1E3A8A", "1E3A8A"]
        
        border_light = {'val': 'single', 'sz': 4, 'color': 'E2E8F0'}
        
        for col_idx, text in enumerate(headers):
            cell_h = temp_table.cell(0, col_idx)
            set_cell_background(cell_h, header_colors[col_idx])
            set_cell_margins(cell_h, top=140, bottom=140, left=150, right=150)
            set_cell_borders(cell_h, top=border_light, bottom=border_light, left=border_light, right=border_light)
            p_h = cell_h.paragraphs[0]
            style_paragraph(p_h, text, font_name="Arial", font_size=9, bold=True, color_rgb=RGBColor(255, 255, 255), align=1 if col_idx > 0 else 0)
            
        phase_inputs = [
            {'name': 'R Phase', 'temp': loc.get('r_temp')},
            {'name': 'Y Phase', 'temp': loc.get('y_temp')},
            {'name': 'B Phase', 'temp': loc.get('b_temp')},
            {'name': 'Neutral', 'temp': loc.get('n_temp')}
        ]
        
        # Compute status for each phase
        def get_phase_status(t_str):
            try:
                t = float(t_str)
                if t > 80: return "Critical (>80°C)", "FEE2E2", "991B1B"
                elif t >= 60: return "Major (60-80°C)", "FFEDD5", "9A3412"
                elif t >= 45: return "Minor (45-60°C)", "FEF9C3", "713F12"
                else: return "Normal (<45°C)", "DCFCE7", "166534"
            except (ValueError, TypeError):
                return "Not Measured", "F1F5F9", "475569"
                
        # Determine highest phase temperature
        max_t = -999.0
        max_idx = -1
        for idx, pi in enumerate(phase_inputs):
            try:
                t_val = float(pi['temp'])
                if t_val > max_t:
                    max_t = t_val
                    max_idx = idx
            except (ValueError, TypeError):
                pass

        for idx, pi in enumerate(phase_inputs):
            row_idx = idx + 1
            cell_name = temp_table.cell(row_idx, 0)
            cell_temp = temp_table.cell(row_idx, 1)
            cell_stat = temp_table.cell(row_idx, 2)
            
            # Zebra striping
            bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
            
            # Apply formatting
            for c in (cell_name, cell_temp, cell_stat):
                set_cell_background(c, bg_color)
                set_cell_margins(c, top=100, bottom=100, left=150, right=150)
                set_cell_borders(c, top=border_light, bottom=border_light, left=border_light, right=border_light)
                
            p_name = cell_name.paragraphs[0]
            style_paragraph(p_name, pi['name'], font_name="Arial", font_size=9.5, bold=True, color_rgb=RGBColor(51, 65, 85))
            
            p_temp = cell_temp.paragraphs[0]
            temp_text = f"{pi['temp']} °C" if pi['temp'] != '' and pi['temp'] is not None else "To be extracted"
            
            # Highlight max temperature phase row
            is_max = (idx == max_idx and max_idx != -1)
            style_paragraph(p_temp, temp_text, font_name="Arial", font_size=9.5, bold=is_max, color_rgb=RGBColor(30, 41, 59) if not is_max else RGBColor(220, 38, 38), align=1)
            
            p_stat = cell_stat.paragraphs[0]
            stat_text, stat_bg, stat_fg = get_phase_status(pi['temp'])
            
            if pi['temp'] != '' and pi['temp'] is not None:
                # Add status badge style inside cell
                set_cell_background(cell_stat, stat_bg)
                fg_r, fg_g, fg_b = int(stat_fg[:2], 16), int(stat_fg[2:4], 16), int(stat_fg[4:], 16)
                style_paragraph(p_stat, stat_text, font_name="Arial", font_size=9, bold=True, color_rgb=RGBColor(fg_r, fg_g, fg_b), align=1)
            else:
                style_paragraph(p_stat, "Pending", font_name="Arial", font_size=9, italic=True, color_rgb=RGBColor(148, 163, 184), align=1)
                
        doc.add_paragraph()
        
        # ----------------------------
        # DIAGNOSTIC CALLOUT CARD
        # ----------------------------
        # Compute overall severity for the callout color border
        overall_priority = 4
        overall_color = "10B981"  # Normal Green
        overall_title = "NORMAL CONDITION"
        
        try:
            temps = [float(t) for t in [loc.get('r_temp'), loc.get('y_temp'), loc.get('b_temp'), loc.get('n_temp')] if t != '' and t is not None]
            if temps:
                max_temp = max(temps)
                if max_temp > 80:
                    overall_priority = 1
                    overall_color = "EF4444" # Red
                    overall_title = "CRITICAL SEVERITY (PRIORITY 1)"
                elif max_temp >= 60:
                    overall_priority = 2
                    overall_color = "F97316" # Orange
                    overall_title = "MAJOR SEVERITY (PRIORITY 2)"
                elif max_temp >= 45:
                    overall_priority = 3
                    overall_color = "EAB308" # Yellow
                    overall_title = "MINOR SEVERITY (PRIORITY 3)"
        except (ValueError, TypeError):
            pass
            
        callout_table = doc.add_table(rows=1, cols=1)
        callout_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        callout_table.autofit = False
        callout_table.columns[0].width = Inches(6.7)
        
        c_cell = callout_table.cell(0, 0)
        set_cell_background(c_cell, "F8FAFC")
        set_cell_margins(c_cell, top=140, bottom=140, left=180, right=150)
        
        # 4pt Left border with corresponding status color, and light grey top/bottom/right borders
        set_cell_borders(c_cell, 
                         left={'val': 'single', 'sz': 24, 'color': overall_color},
                         top={'val': 'single', 'sz': 4, 'color': 'E2E8F0'},
                         bottom={'val': 'single', 'sz': 4, 'color': 'E2E8F0'},
                         right={'val': 'single', 'sz': 4, 'color': 'E2E8F0'})
        
        p_c_title = c_cell.paragraphs[0]
        p_c_title.paragraph_format.space_after = Pt(4)
        c_r, c_g, c_b = int(overall_color[:2], 16), int(overall_color[2:4], 16), int(overall_color[4:], 16)
        style_paragraph(p_c_title, f"DIAGNOSIS: {overall_title}", font_name="Arial", font_size=10, bold=True, color_rgb=RGBColor(c_r, c_g, c_b))
        
        p_obs = c_cell.add_paragraph()
        p_obs.paragraph_format.space_after = Pt(2)
        r_obs = p_obs.add_run("Observation: ")
        r_obs.bold = True
        r_obs.font.name = 'Arial'
        r_obs.font.size = Pt(9.5)
        r_obs.font.color.rgb = RGBColor(71, 85, 105)
        
        obs_text = loc.get('observation', 'Temperature found to be Normal.')
        r_obs_val = p_obs.add_run(obs_text)
        r_obs_val.font.name = 'Arial'
        r_obs_val.font.size = Pt(9.5)
        r_obs_val.font.color.rgb = RGBColor(15, 23, 42)
        
        p_rem = c_cell.add_paragraph()
        p_rem.paragraph_format.space_after = Pt(0)
        r_rem = p_rem.add_run("Recommendation / Remarks: ")
        r_rem.bold = True
        r_rem.font.name = 'Arial'
        r_rem.font.size = Pt(9.5)
        r_rem.font.color.rgb = RGBColor(71, 85, 105)
        
        rem_text = loc.get('remarks', 'Working in Normal Condition.')
        r_rem_val = p_rem.add_run(rem_text)
        r_rem_val.font.name = 'Arial'
        r_rem_val.font.size = Pt(9.5)
        r_rem_val.font.color.rgb = RGBColor(15, 23, 42)
        
        # Page break if not the last item
        if index < len(locations) - 1:
            doc.add_page_break()
            
    # Save document
    report_filename = "Thermography_Report.docx"
    report_path = os.path.join(session_report_dir, report_filename)
    doc.save(report_path)
    
    return jsonify({
        'download_url': f"/api/download/{session_id}/{report_filename}"
    })


@app.route('/api/download/<session_id>/<filename>')
def download_report_file(session_id, filename):
    session_report_dir = os.path.join(app.config['REPORT_FOLDER'], session_id)
    return send_from_directory(session_report_dir, filename, as_attachment=True)


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
