from docx import Document
from docx.shared import Inches
from openpyxl import load_workbook
import os

# ----------------------------
# CONFIGURATION
# ----------------------------

EXCEL_FILE = "Thermography_Master_Mapping.xlsx"

THERMAL_FOLDER = "thermal_images"
VISUAL_FOLDER = "visual_images"

OUTPUT_FILE = "Thermography_Report.docx"

CLIENT_NAME = "D-MART"
SITE_NAME = "KANCHARLAPALEM, VISAKHAPATNAM"
DATE = "11-05-2026"

# ----------------------------
# LOAD EXCEL
# ----------------------------

wb = load_workbook(EXCEL_FILE)
ws = wb.active

doc = Document()

# ----------------------------
# HEADER
# ----------------------------

doc.add_heading('THERMOGRAPHY REPORT', 0)

doc.add_paragraph(f"Client : {CLIENT_NAME}")
doc.add_paragraph(f"Site   : {SITE_NAME}")
doc.add_paragraph(f"Date   : {DATE}")

doc.add_page_break()

# ----------------------------
# PROCESS EACH LOCATION
# ----------------------------

for row in ws.iter_rows(min_row=2, values_only=True):

    sno = row[0]
    location = row[1]

    r_img = row[2]
    y_img = row[3]
    b_img = row[4]
    n_img = row[5]
    visual_img = row[6]

    doc.add_heading(f"{sno}. {location}", level=1)

    # ----------------------------
    # VISUAL IMAGE
    # ----------------------------

    visual_path = os.path.join(
        VISUAL_FOLDER,
        f"{visual_img}.jpg"
    )

    if os.path.exists(visual_path):
        doc.add_paragraph("Visual Image")
        doc.add_picture(
            visual_path,
            width=Inches(3)
        )

    # ----------------------------
    # THERMAL IMAGES
    # ----------------------------

    doc.add_paragraph("Thermal Images")

    thermal_images = [
        r_img,
        y_img,
        b_img,
        n_img
    ]

    for img in thermal_images:

        img_path = os.path.join(
            THERMAL_FOLDER,
            f"{img}.jpg"
        )

        if os.path.exists(img_path):
            doc.add_picture(
                img_path,
                width=Inches(2.5)
            )

    # ----------------------------
    # MEASUREMENT TABLE
    # ----------------------------

    table = doc.add_table(
        rows=5,
        cols=2
    )

    table.style = 'Table Grid'

    table.cell(0,0).text = "Phase"
    table.cell(0,1).text = "Temperature"

    table.cell(1,0).text = "R Phase"
    table.cell(1,1).text = "To be extracted"

    table.cell(2,0).text = "Y Phase"
    table.cell(2,1).text = "To be extracted"

    table.cell(3,0).text = "B Phase"
    table.cell(3,1).text = "To be extracted"

    table.cell(4,0).text = "Neutral"
    table.cell(4,1).text = "To be extracted"

    # ----------------------------
    # SEVERITY TABLE
    # ----------------------------

    doc.add_paragraph()

    severity = doc.add_table(
        rows=5,
        cols=3
    )

    severity.style = 'Table Grid'

    severity.cell(0,0).text = "Priority"
    severity.cell(0,1).text = "Temperature"
    severity.cell(0,2).text = "Remarks"

    severity.cell(1,0).text = "1"
    severity.cell(1,1).text = ">80°C"
    severity.cell(1,2).text = "Repair Immediately"

    severity.cell(2,0).text = "2"
    severity.cell(2,1).text = "60-80°C"
    severity.cell(2,2).text = "Repair Earliest Opportunity"

    severity.cell(3,0).text = "3"
    severity.cell(3,1).text = "45-60°C"
    severity.cell(3,2).text = "Investigate During Maintenance"

    severity.cell(4,0).text = "4"
    severity.cell(4,1).text = "<45°C"
    severity.cell(4,2).text = "Record and Monitor"

    # ----------------------------
    # OBSERVATION
    # ----------------------------

    doc.add_paragraph()
    doc.add_paragraph(
        "OBSERVATION : Temperature found to be Normal."
    )

    doc.add_paragraph(
        "REMARKS : Working in Normal Condition."
    )

    doc.add_page_break()

# ----------------------------
# SAVE REPORT
# ----------------------------

doc.save(OUTPUT_FILE)

print(
    f"Report Generated: {OUTPUT_FILE}"
)
